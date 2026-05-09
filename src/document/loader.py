"""
文档加载器
支持多种文档格式的加载
"""

from typing import List
from pathlib import Path
from loguru import logger

from langchain.schema import Document
import fitz  # PyMuPDF
from docx import Document as DocxDocument


class DocumentLoader:
    """多格式文档加载器"""
    
    def __init__(self):
        """初始化文档加载器"""
        self.supported_formats = {
            ".pdf": self._load_pdf,
            ".docx": self._load_docx,
            ".doc": self._load_docx,
            ".txt": self._load_txt,
            ".md": self._load_markdown
        }
        
        logger.info(f"文档加载器初始化完成，支持格式: {list(self.supported_formats.keys())}")
    
    def load(self, file_path: str) -> List[Document]:
        """
        加载文档
        
        Args:
            file_path: 文档路径
            
        Returns:
            文档列表
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {extension}")
        
        loader_func = self.supported_formats[extension]
        documents = loader_func(file_path)
        
        # 添加文件元数据
        for doc in documents:
            doc.metadata.setdefault("source", file_path)
            doc.metadata.setdefault("file_name", path.name)
            doc.metadata.setdefault("file_type", extension)
        
        logger.info(f"成功加载文档: {file_path}")
        return documents
    
    def _load_pdf(self, file_path: str) -> List[Document]:
        """加载PDF文档"""
        documents = []
        
        with fitz.open(file_path) as pdf:
            for page_num, page in enumerate(pdf, 1):
                text = page.get_text()
                
                if text.strip():  # 忽略空白页
                    doc = Document(
                        page_content=text,
                        metadata={
                            "page": page_num,
                            "total_pages": len(pdf)
                        }
                    )
                    documents.append(doc)
        
        return documents
    
    def _load_docx(self, file_path: str) -> List[Document]:
        """加载Word文档"""
        documents = []
        
        doc = DocxDocument(file_path)
        
        # 按段落组织内容
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # 将所有段落合并为一个文档
        if paragraphs:
            full_text = "\n\n".join(paragraphs)
            documents.append(Document(
                page_content=full_text,
                metadata={"paragraphs": len(paragraphs)}
            ))
        
        return documents
    
    def _load_txt(self, file_path: str) -> List[Document]:
        """加载TXT文档"""
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        
        return [Document(page_content=text, metadata={})]
    
    def _load_markdown(self, file_path: str) -> List[Document]:
        """加载Markdown文档"""
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
        
        # 可以在这里添加Markdown特定的处理逻辑
        # 例如：提取标题、代码块等
        
        return [Document(page_content=text, metadata={})]
    
    def load_batch(self, file_paths: List[str]) -> List[Document]:
        """
        批量加载文档
        
        Args:
            file_paths: 文档路径列表
            
        Returns:
            所有文档列表
        """
        all_documents = []
        
        for file_path in file_paths:
            try:
                documents = self.load(file_path)
                all_documents.extend(documents)
            except Exception as e:
                logger.error(f"加载文档失败 {file_path}: {e}")
        
        return all_documents
